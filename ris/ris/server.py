from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os, uuid, datetime
from docx import Document
import mammoth

app = Flask(__name__)
CORS(app)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
STORAGE_DIR = os.path.join(BASE_DIR, "storage", "reports")
os.makedirs(STORAGE_DIR, exist_ok=True)

# in-memory index (replace with DB later if needed)
REPORT_INDEX = {}

@app.route("/")
def home():
    return "Malwa RIS backend running"

# ---------------- WORD → HTML ----------------
@app.route("/api/word_to_html", methods=["POST"])
def word_to_html():
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No file"}), 400

    with mammoth.convert_to_html(file) as result:
        html = result.value

    return jsonify({"html": html})

# ---------------- SAVE REPORT ----------------
@app.route("/api/save_report", methods=["POST"])
def save_report():
    data = request.json
    report_id = str(uuid.uuid4())

    filename = f"{data['patient_name']}_{data['uhid']}_{report_id}.docx".replace(" ", "_")
    path = os.path.join(STORAGE_DIR, filename)

    doc = Document()
    doc.add_paragraph(
        f"Patient Name: {data['patient_name']}    "
        f"UHID: {data['uhid']}    "
        f"Modality: {data['modality']}"
    )
    doc.add_paragraph("")

    # VERY IMPORTANT: paragraph-style, not line-by-line garbage
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(data["html"], "html.parser")
    for p in soup.find_all(["p", "li"]):
        text = p.get_text(" ").strip()
        if text:
            doc.add_paragraph(text)

    doc.save(path)

    REPORT_INDEX[report_id] = {
        "id": report_id,
        "patient": data["patient_name"],
        "uhid": data["uhid"],
        "modality": data["modality"],
        "filename": filename,
        "date": datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    }

    return jsonify({"id": report_id})

# ---------------- FETCH REPORTS ----------------
@app.route("/api/reports")
def fetch_reports():
    return jsonify(list(REPORT_INDEX.values()))

# ---------------- DOWNLOAD WORD ----------------
@app.route("/api/download_word/<rid>")
def download_word(rid):
    report = REPORT_INDEX.get(rid)
    if not report:
        return "Not found", 404

    path = os.path.join(STORAGE_DIR, report["filename"])
    if not os.path.exists(path):
        return "File missing", 404

    return send_file(path, as_attachment=True)

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)








